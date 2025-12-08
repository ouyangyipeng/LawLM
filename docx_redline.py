from pathlib import Path
from typing import Iterable, List

from docx import Document
from docx.shared import RGBColor
from docx.text.run import Run


def _add_diff_runs(paragraph, tokens: List[str], diff_flags: List[str]) -> None:
    """Add runs to paragraph with coloring/strike/underline based on diff flags.

    diff_flags values:
    - 'equal': normal text
    - 'insert': green + underline
    - 'delete': red + strikethrough
    """
    for token, flag in zip(tokens, diff_flags):
        run: Run = paragraph.add_run(token)
        if flag == "insert":
            run.font.color.rgb = RGBColor(0, 128, 0)
            run.font.underline = True
        elif flag == "delete":
            run.font.color.rgb = RGBColor(220, 0, 0)
            run.font.strike = True
        # add trailing space to separate tokens visually
        paragraph.add_run(" ")


def build_redline_docx(
    output_path: Path,
    ocr_text: str,
    corrected_text: str,
    risks: str,
    safety_score: float,
    similarity: float,
) -> None:
    """Build a Word document with redline-like markup and risk comments section."""
    doc = Document()
    doc.add_heading("法律文书自动修订稿", level=1)

    doc.add_heading("修订概览", level=2)
    doc.add_paragraph(f"安全评分: {safety_score}")
    doc.add_paragraph(f"文本相似度(自动评测): {similarity:.3f}")

    doc.add_heading("红线修订", level=2)
    redline_para = doc.add_paragraph()
    tokens, flags = _word_level_diff(ocr_text, corrected_text)
    _add_diff_runs(redline_para, tokens, flags)

    doc.add_heading("纠正后全文", level=2)
    doc.add_paragraph(corrected_text or "(empty)")

    doc.add_heading("风险批注", level=2)
    if risks:
        for idx, item in enumerate(_split_risks(risks), 1):
            doc.add_paragraph(f"[{idx}] {item}")
    else:
        doc.add_paragraph("未检测到风险。")

    doc.save(output_path)


def _word_level_diff(a: str, b: str) -> (List[str], List[str]):
    """Compute a simple word-level diff returning tokens and flags."""
    import difflib

    a_tokens = a.split()
    b_tokens = b.split()
    matcher = difflib.SequenceMatcher(a=a_tokens, b=b_tokens)

    tokens: List[str] = []
    flags: List[str] = []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            tokens.extend(a_tokens[i1:i2])
            flags.extend(["equal"] * (i2 - i1))
        elif tag == "delete":
            tokens.extend(a_tokens[i1:i2])
            flags.extend(["delete"] * (i2 - i1))
        elif tag == "insert":
            tokens.extend(b_tokens[j1:j2])
            flags.extend(["insert"] * (j2 - j1))
        elif tag == "replace":
            tokens.extend(a_tokens[i1:i2])
            flags.extend(["delete"] * (i2 - i1))
            tokens.extend(b_tokens[j1:j2])
            flags.extend(["insert"] * (j2 - j1))
    return tokens, flags


def _split_risks(risks) -> List[str]:
    # Accept str or list and split by common delimiters.
    if not risks:
        return []
    if isinstance(risks, list):
        parts = risks
    else:
        parts = risks.replace("\n", ";").split(";")
    cleaned = [p.strip() for p in parts if p and str(p).strip()]
    return cleaned
